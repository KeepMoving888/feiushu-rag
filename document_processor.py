"""
文档处理模块
实现：
1) 文档解析：PDF / Word / Excel / TXT
2) 文本分块
"""

from __future__ import annotations

import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


class DocumentProcessorError(Exception):
    """文档处理异常。"""


@dataclass
class ParsedDocument:
    """解析后的文档结构。"""

    text: str
    metadata: dict[str, Any]


class DocumentProcessor:
    """文档解析与分块。"""

    def __init__(self, chunk_size: int | None = None, chunk_overlap: int | None = None) -> None:
        """初始化分块器。

        参数优先级：
        1) 显式传入参数
        2) .env / 环境变量
           - DOC_CHUNK_SIZE（默认 700）
           - DOC_CHUNK_OVERLAP（默认 120）
        """

        env_chunk_size = os.getenv("DOC_CHUNK_SIZE", "700")
        env_chunk_overlap = os.getenv("DOC_CHUNK_OVERLAP", "120")

        try:
            final_chunk_size = int(chunk_size if chunk_size is not None else env_chunk_size)
        except (TypeError, ValueError):
            final_chunk_size = 700

        try:
            final_chunk_overlap = int(chunk_overlap if chunk_overlap is not None else env_chunk_overlap)
        except (TypeError, ValueError):
            final_chunk_overlap = 120

        # 基础容错，防止 overlap >= size 导致切分异常
        if final_chunk_size <= 0:
            final_chunk_size = 700
        if final_chunk_overlap < 0:
            final_chunk_overlap = 0
        if final_chunk_overlap >= final_chunk_size:
            final_chunk_overlap = max(0, final_chunk_size // 5)

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=final_chunk_size,
            chunk_overlap=final_chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", " "],
        )

    def parse_file(self, file_path: str) -> ParsedDocument:
        """解析文件内容。"""

        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise DocumentProcessorError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                text = self._parse_pdf(path)
            elif suffix in {".docx", ".doc"}:
                text = self._parse_word(path)
            elif suffix in {".xlsx", ".xls"}:
                text = self._parse_excel(path)
            elif suffix == ".txt":
                text = self._parse_txt(path)
            else:
                raise DocumentProcessorError(f"不支持的文件格式: {suffix}")
        except Exception as exc:
            if isinstance(exc, DocumentProcessorError):
                raise
            raise DocumentProcessorError(f"解析文件失败: {exc}") from exc

        cleaned = text.strip()
        if not cleaned:
            raise DocumentProcessorError("文档内容为空或无法解析出文本")

        return ParsedDocument(
            text=cleaned,
            metadata={
                "source": str(path),
                "filename": path.name,
                "suffix": suffix,
            },
        )

    def split_text(self, text: str, base_metadata: dict[str, Any] | None = None) -> list[dict[str, Any]]:
        """文本分块并附带 metadata。"""

        if not text or not text.strip():
            raise DocumentProcessorError("待分块文本不能为空")

        chunks = self.splitter.split_text(text)
        base_metadata = base_metadata or {}
        result = []
        for idx, chunk in enumerate(chunks):
            result.append(
                {
                    "chunk_id": idx,
                    "text": chunk,
                    "metadata": {**base_metadata, "chunk_index": idx},
                }
            )
        return result

    def process_bitable_data(
        self,
        headers: list[str] | dict[str, str],
        records: list[dict[str, Any]],
        app_token: str,
        table_id: str,
        table_name: str,
    ) -> tuple[list[str], list[dict[str, Any]]]:
        """处理飞书多维表格数据（行级分块）。

        设计说明：
        1) 每一行记录生成一个独立文本块（行级 chunk）。
        2) 自动过滤空行、无效行。
        3) 对日期、数字、人员、附件等常见复杂字段做可读化转换。

        参数：
        - headers: 表头信息。支持两种格式：
          a) list[str]：字段名列表
          b) dict[str, str]：字段名 -> 字段ID（或反向映射）
        - records: 多维表格记录列表（每条通常含 record_id + fields）
        - app_token / table_id / table_name: 表标识信息

        返回：
        - texts: 处理后的文本块列表
        - metadatas: 与文本块一一对应的 metadata 列表
        """

        if not app_token:
            raise DocumentProcessorError("app_token 不能为空")
        if not table_id:
            raise DocumentProcessorError("table_id 不能为空")
        if not table_name or not table_name.strip():
            table_name = table_id

        # 统一字段名列表
        header_names: list[str] = []
        if isinstance(headers, list):
            header_names = [str(h).strip() for h in headers if str(h).strip()]
        elif isinstance(headers, dict):
            # 优先使用“字段名 -> 字段ID”场景；若是“字段ID -> 字段名”也能兼容
            header_names = [str(k).strip() for k in headers.keys() if str(k).strip()]
        else:
            raise DocumentProcessorError("headers 仅支持 list[str] 或 dict[str, str]")

        if not isinstance(records, list):
            raise DocumentProcessorError("records 必须是 list")

        texts: list[str] = []
        metadatas: list[dict[str, Any]] = []
        header_text = ", ".join(header_names)

        for row_index, rec in enumerate(records):
            if not isinstance(rec, dict):
                continue

            record_id = str(rec.get("record_id") or rec.get("id") or "").strip()
            fields = rec.get("fields", {})
            if not isinstance(fields, dict):
                fields = {}

            row_pairs: list[str] = []

            # 若提供了表头，优先按表头顺序输出，保持可读性
            if header_names:
                for h in header_names:
                    value = fields.get(h)
                    text_value = self._format_bitable_value(value)
                    if text_value:
                        row_pairs.append(f"{h}: {text_value}")
            else:
                for k, v in fields.items():
                    key = str(k).strip()
                    if not key:
                        continue
                    text_value = self._format_bitable_value(v)
                    if text_value:
                        row_pairs.append(f"{key}: {text_value}")

            # 过滤空行：没有有效字段值则跳过
            if not row_pairs:
                continue

            row_data_text = ", ".join(row_pairs)
            chunk_text = (
                f"表格名称：{table_name}\n"
                f"表头：{header_text}\n"
                f"该行数据：{row_data_text}"
            )

            # 再做一次空文本防御
            if not chunk_text.strip() or "该行数据：" == chunk_text.strip().split("\n")[-1]:
                continue

            metadata = {
                "source_type": "bitable",
                "app_token": app_token,
                "table_id": table_id,
                "record_id": record_id or f"row_{row_index}",
                "table_name": table_name,
                "row_index": row_index,
            }

            texts.append(chunk_text)
            metadatas.append(metadata)

        return texts, metadatas

    def _format_bitable_value(self, value: Any) -> str:
        """将多维表格字段值转换为可读文本。"""

        if value is None:
            return ""

        # 简单类型
        if isinstance(value, (str, int, float, bool)):
            text = str(value).strip()
            return "" if text in {"", "None", "nan", "NaN"} else text

        # list: 常见于多选、人员列表、附件列表
        if isinstance(value, list):
            parts: list[str] = []
            for item in value:
                part = self._format_bitable_value(item)
                if part:
                    parts.append(part)
            return ", ".join(parts)

        # dict: 常见于人员、日期、附件、关联记录等结构
        if isinstance(value, dict):
            # 人员字段常见键
            for key in ("name", "en_name", "display_name", "full_name"):
                name = value.get(key)
                if isinstance(name, str) and name.strip():
                    return name.strip()

            # 日期字段常见键（时间戳）
            if "timestamp" in value:
                ts = value.get("timestamp")
                if isinstance(ts, (int, float)):
                    # 飞书常见毫秒时间戳
                    sec = ts / 1000 if ts > 10_000_000_000 else ts
                    try:
                        return str(pd.to_datetime(sec, unit="s"))
                    except Exception:
                        return str(ts)

            # 常见文本键兜底
            for key in ("text", "value", "url", "id"):
                v = value.get(key)
                if isinstance(v, (str, int, float)) and str(v).strip():
                    return str(v).strip()

            # 最终兜底 JSON 文本
            return str(value)

        # 其他对象兜底
        return str(value).strip()

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        reader = PdfReader(str(path))
        texts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            texts.append(page_text)
        return "\n".join(texts)

    @staticmethod
    def _parse_word(path: Path) -> str:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paragraphs)

    @staticmethod
    def _parse_excel(path: Path) -> str:
        """解析 Excel：按 sheet 拼接文本。"""

        excel = pd.read_excel(path, sheet_name=None)
        blocks: list[str] = []
        for sheet_name, df in excel.items():
            if df.empty:
                continue
            csv_text = df.fillna("").astype(str).to_csv(index=False)
            blocks.append(f"[Sheet: {sheet_name}]\n{csv_text}")
        return "\n\n".join(blocks)

    @staticmethod
    def _parse_txt(path: Path) -> str:
        # 先尝试 UTF-8，失败后降级 GBK
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="gbk", errors="ignore")
